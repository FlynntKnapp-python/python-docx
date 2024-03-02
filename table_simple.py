import os

from docx import Document

# Specify the file path for the .docx file
file_path = "output/TableSimple.docx"

# Delete the file if it exists
# Check if the file exists to avoid FileNotFoundError
if os.path.exists(file_path):
    # Delete the file
    os.remove(file_path)
    print(f"File {file_path} deleted.")
else:
    print("The file does not exist.")

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading("The Official Table Example!", 0)

records = (
    ("Skill 00", "Skill 01"),
    ("Skill 10", "Skill 11"),
    ("Skill 20", "Skill 21"),
)

# Add a table to the document
table = doc.add_table(rows=1, cols=2)
for skill0, skill1 in records:
    row_cells = table.add_row().cells
    row_cells[0].text = skill0
    row_cells[1].text = skill1

# Save the document to a .docx file
doc.save(file_path)
