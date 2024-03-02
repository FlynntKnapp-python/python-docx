from docx import Document

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading("Document Title", 0)

# Add a paragraph of text
p = doc.add_paragraph(
    "This is a simple paragraph that is being added to the document. "
)

# Save the document to a .docx file
doc.save("output/HeadingAndParagraph.docx")
