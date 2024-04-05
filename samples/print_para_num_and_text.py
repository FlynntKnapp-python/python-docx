# samples\print_para_num_and_text.py

from docx import Document

# Use `enumerate_paragraphs` to print the number and text of each paragraph in the document:
from utils import enumerate_paragraphs

# Specify the file path for the .docx file:
file_path = "samples/output/ParaTablePara.docx"

# Load the .docx file (instantiating the Document object):
doc = Document(file_path)

# Get the number of paragraphs in the document:
num_paragraphs = len(doc.paragraphs)

# Print the number of paragraphs in the document:
print(f"Number of Paragraphs: {num_paragraphs}")

# Print the text of each paragraph in the document:
for i, paragraph in enumerate(doc.paragraphs):
    print(f"Paragraph {i}: {paragraph.text}")

list_result = enumerate_paragraphs(doc)
print("list_result: ", list_result)
print("type(list_result): ", type(list_result))
print("type(list_result[0]): ", type(list_result[0]))


for item in list_result:
    print(item)

for item in list_result:
    print(item[0], item[1])