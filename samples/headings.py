from docx import Document
from utils import create_docx_if_not_exists, delete_and_or_save_docx

# Specify the file path for the .docx file
file_path = "samples/output/Headings.docx"

# Create a new Document
doc = Document()

# Add a title to the document
h0 = doc.add_heading("Heading 0", 0)
print("Added h.text: \n", h0.text)

# Add a smaller heading to the document
h1 = doc.add_heading("Heading 1", level=1)
print("Added h.text: \n", h1.text)

# Add a smaller heading to the document
h2 = doc.add_heading("Heading 2", level=2)
print("Added h.text: \n", h2.text)

# Add a smaller heading to the document
h3 = doc.add_heading("Heading 3", level=3)
print("Added h.text: \n", h3.text)

# Save the document to a .docx file
delete_and_or_save_docx(file_path, doc)
