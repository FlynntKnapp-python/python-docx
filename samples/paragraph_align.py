# samples\paragraph_align.py

from base import docx_builder
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Specify the file path for the .docx file:
file_path = "samples/output/ParagraphAlign.docx"

# Create a new Document:
doc = Document()

# Add a title to the document:
doc.add_heading("Paragraph Alignment", 0)

# Add some paragraphs to the document:
paragraph_left = doc.add_paragraph("This paragraph is left-aligned.")
paragraph_right = doc.add_paragraph("This paragraph is right-aligned.")

# Set the alignment of the paragraphs:
paragraph_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Save the document to a .docx file:
saved = docx_builder.manage_docx_file(file_path, doc, "save")
print("Saved: ", saved)
