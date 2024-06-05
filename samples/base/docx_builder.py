# samples\base\docx_builder.py

import math
import os
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


def manage_docx_file(
    path: str, document: Document = None, action: str = "load_or_create"
) -> Document:
    """
    Manage a .docx file (load, create, or delete and save).

    Parameters:
    - path (str): The file system path where the .docx file is saved.
    - document (Document): The docx.Document object. Required for 'delete_and_create'
    and 'save' action.
    - action (str): Action to perform on the file: 'load_or_create',
    'delete_and_create', or 'save'. Default is 'load_or_create'.

    Returns:
    - Document: For 'load_or_create' and 'delete_and_create' actions, returns a
        docx.Document object.
    - bool: For 'save' action, returns True if the document was saved successfully,
    False otherwise.
    """
    file_exists = os.path.exists(path)

    if action == "load_or_create":
        if file_exists:
            print(f"Loading {path}...")
            return Document(path)
        else:
            print(f"The file {path} does not exist. Creating a new one...")
            return Document()

    if action == "delete_and_create":
        if file_exists:
            print(f"Deleting {path}...")
            os.remove(path)
        print(f"Creating {path}...")
        return Document()

    if action == "save":
        try:
            print(f"Saving the document to {path}...")
            document.save(path)
            print(f"Document saved to {path}.")
            return True
        except Exception as e:
            print(f"Error saving the document: {e}")
            return False


def save_docx(path: str, document: Document):
    """
    Save a document to a .docx file.

    Parameters:
    - path (str): The file system path where the .docx file is saved.
    - document (Document): The docx.Document object to be saved.

    Returns:
    - bool: True if the document was saved successfully, False otherwise.
    """
    try:
        print(f"Saving the document to {path}...")
        document.save(path)
        print(f"Document saved to {path}.")
        return True
    except Exception as e:
        print(f"Error saving the document: {e}")
        return False


def add_table(doc: Document, items: list, cols: int) -> Document:
    """
    Create and populate a table in the document.

    Parameters:
    - doc (Document): The Document object to add the table to.
    - items (list): List of items to populate the table.
    - cols (int): Number of columns in the table.

    Returns:
    - Document: The modified Document object.
    """
    rows = (len(items) + cols - 1) // cols  # Calculate rows needed
    table = doc.add_table(rows=rows, cols=cols)

    for i, item in enumerate(items):
        row, col = divmod(i, cols)
        table.cell(row, col).text = str(item)
        for paragraph in table.cell(row, col).paragraphs:
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0

    return doc


def add_resume_heading(doc: Document, name: str, title: str) -> Document:
    """
    Add a heading to a resume document.

    Parameters:
    - doc (Document): The Document object to add the heading to.
    - name (str): The name to add to the heading.
    - title (str): The title to add to the heading.

    Returns:
    - Document: The modified Document object.
    """
    name_paragraph = doc.add_paragraph()
    name_paragraph.paragraph_format.space_after = 0
    # Alternatively, use the following to set the space after:
    # name_paragraph.paragraph_format.space_after = Pt(<NN>)
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_run = name_paragraph.add_run(name)
    name_run.font.name = "Times New Roman"
    name_run.font.size = Pt(18)
    name_run.bold = True
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(12)

    return doc


def add_resume_heading_as_table(doc: Document, name: str, title: str) -> Document:
    """
    Add a heading to a resume document as a table.

    Parameters:
    - doc (Document): The Document object to add the heading to.
    - name (str): The name to add to the heading.
    - title (str): The title to add to the heading.

    Returns:
    - Document: The modified Document object.
    """
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = name
    table.cell(1, 0).text = title

    return doc


def add_resume_address(
    doc: Document, address: str, city: str, state: str, zip: str
) -> Document:
    """
    Add an address to a resume document.

    Parameters:
    - doc (Document): The Document object to add the address to.
    - address (str): The address to add to the document.
    - city (str): The city to add to the document.
    - state (str): The state to add to the document.
    - zip (str): The ZIP code to add to the document.

    Returns:
    - Document: The modified Document object.
    """
    address_paragraph = doc.add_paragraph()
    address_paragraph.paragraph_format.space_after = 0
    address_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    address_run = address_paragraph.add_run(address)
    address_run.font.name = "Times New Roman"
    address_run.font.size = Pt(12)
    city_state_zip_paragraph = doc.add_paragraph()
    city_state_zip_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    city_state_zip_run = city_state_zip_paragraph.add_run(f"{city}, {state} {zip}")
    city_state_zip_run.font.name = "Times New Roman"
    city_state_zip_run.font.size = Pt(12)

    return doc


def add_resume_phone(doc: Document, phone: str) -> Document:
    """
    Add a phone number to a resume document.

    Parameters:
    - doc (Document): The Document object to add the phone number to.
    - phone (str): The phone number to add to the document.

    Returns:
    - Document: The modified Document object.
    """
    phone_paragraph = doc.add_paragraph()
    phone_paragraph.paragraph_format.space_after = 0
    phone_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    phone_run = phone_paragraph.add_run(f"Phone: {phone}")
    phone_run.font.name = "Times New Roman"
    phone_run.font.size = Pt(12)

    return doc


def add_resume_personal_links(doc: Document, links: dict) -> Document:
    """
    Add personal links to a resume document.

    Parameters:
    - doc (Document): The Document object to add the links to.
    - links (dict): A dictionary of personal links to add to the document.

    Returns:
    - Document: The modified Document object.
    """
    for key, value in links.items():
        link_paragraph = doc.add_paragraph()
        link_paragraph.paragraph_format.space_before = 57150
        link_paragraph.paragraph_format.space_after = 57150
        link_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        link_run = link_paragraph.add_run(f"{key}: {value}")
        link_run.font.name = "Times New Roman"
        link_run.font.size = Pt(12)

    return doc


def set_margins(
    doc: Document,
    top: float = 1.0,
    bottom: float = 1.0,
    left: float = 1.0,
    right: float = 1.0,
) -> Document:
    """
    Set the margins of a Word document.

    Parameters:
    - doc (Document): The Document object to set the margins for.
    - top (int): The top margin in inches.
    - bottom (int): The bottom margin in inches.
    - left (int): The left margin in inches.
    - right (int): The right margin in inches.

    Returns:
    - Document: The modified Document object.

    Conversion Factors:
    0.0625 Inches - 57150
    0.125 Inches - 114300
    0.25 Inches - 228600
    0.5 Inches - 457200
    0.75 Inches - 685800
    1.0 Inches - 914400
    1.25 Inches - 1143000
    1.5 Inches - 1371600
    """
    conversion_factor = 914400

    section = doc.sections[0]
    section.top_margin = math.floor(top * conversion_factor)
    section.bottom_margin = math.floor(bottom * conversion_factor)
    section.left_margin = math.floor(left * conversion_factor)
    section.right_margin = math.floor(right * conversion_factor)

    return doc


def insert_horizontal_line_paragraph_bottom(
    paragraph: Paragraph, thickness: int = 1
) -> Paragraph:
    """
    Insert a horizontal line into a Word document paragraph. The paragraph is expected
    to be an instance from a list of paragraphs obtained via `Document.paragraphs` from
    the python-docx library.

    Parameters:
    - paragraph (Paragraph): The paragraph instance to insert the line into. This should
                              be obtained from the list of paragraphs in a Document
                              object (`Document.paragraphs`).
    - thickness (int): The width of the line in points, with a default of 1 points.

    Returns:
    - Paragraph: The modified paragraph with a horizontal line added.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom_bdr = OxmlElement("w:bottom")
    bottom_bdr.set(qn("w:val"), "single")
    bottom_bdr.set(qn("w:sz"), str(thickness * 8))  # Size of the border, e.g., 4pt
    bottom_bdr.set(qn("w:space"), "1")
    bottom_bdr.set(qn("w:color"), "auto")
    pBdr.append(bottom_bdr)
    pPr.append(pBdr)
    return paragraph


def insert_horizontal_line_paragraph_top(
    paragraph: Paragraph, thickness: int = 1
) -> Paragraph:
    """
    Insert a horizontal line at the top of a Word document paragraph. The paragraph is
    expected to be an instance from a list of paragraphs obtained via
    `Document.paragraphs` from the python-docx library.

    Parameters:
    - paragraph (Paragraph): The paragraph instance to insert the line into. This should
                              be obtained from the list of paragraphs in a Document
                              object (`Document.paragraphs`).
    - thickness (int): The width of the line in points, with a default of 1 points.

    Returns:
    - Paragraph: The modified paragraph with a horizontal line added at the top.
    """
    paragraph.paragraph_format.space_before = 114300
    pPr = paragraph._p.get_or_add_pPr()  # Get or add paragraph properties
    pBdr = OxmlElement("w:pBdr")  # Create a paragraph border element
    top_bdr = OxmlElement("w:top")  # Create a top border element
    top_bdr.set(qn("w:val"), "single")  # Border type, e.g., single line
    # Size of the border, e.g., 4pt (note: Word's measurement unit is 1/8th of a point)
    top_bdr.set(qn("w:sz"), str(thickness * 8))
    top_bdr.set(qn("w:space"), "1")  # The space between the border and the text
    top_bdr.set(qn("w:color"), "auto")  # Border color
    pBdr.append(top_bdr)  # Append the top border to the paragraph border
    pPr.append(pBdr)  # Append the paragraph border to the paragraph properties
    return paragraph


def insert_horizontal_line(doc: Document, thickness: int = 1) -> Document:
    """
    Insert a horizontal line into a Word document. This will create an empty paragraph
    with a bottom border.

    Parameters:
    - doc (Document): The Document object to insert the line into.
    - thickness (int): The width of the line in points.

    Returns:
    - Document: The modified Document object.
    """
    paragraph = doc.add_paragraph()
    # paragraph.paragraph_format.space_before = 0
    # paragraph.paragraph_format.space_after = 0
    # paragraph.paragraph_format.line_spacing = 0
    paragraph.add_run(" ").font.size = Pt(1)
    paragraph = insert_horizontal_line_paragraph_top(paragraph, thickness)
    return doc


def enumerate_paragraphs(doc: Document) -> list:
    """
    Returns an enumerated list of paragraphs in a Word document.

    Parameters:
    - doc (Document): The Document object.

    Returns:
    - list: The enumerated list of paragraphs in the document.
    """
    return [(i, paragraph.text) for i, paragraph in enumerate(doc.paragraphs)]


def print_attributes(obj: Any) -> None:
    """
    Print the attributes of an object to the console.

    Parameters:
    - obj: The object for which to list attributes.

    Returns:
    - None
    """
    print(f"List of {type(obj).__name__} attributes:")
    for attr in dir(obj):
        print(f"\t{attr}")

    return None


def list_paragraphs(doc: Document) -> None:
    """
    Print the paragraphs of a Word document to the console.

    Parameters:
    - doc (Document): The Document object.

    Returns:
    - None
    """
    if len(doc.paragraphs) == 0:
        print("The document does not contain any root paragraphs.")
    else:
        print(f"The document contains {len(doc.paragraphs)} root paragraphs.")
        for i, paragraph in enumerate(doc.paragraphs):
            print(f"\nParagraph ({i}):")
            print(paragraph.text)

    return None


def list_tables(doc: Document) -> None:
    """
    Print the tables of a Word document to the console.

    Parameters:
    - doc (Document): The Document object.

    Returns:
    - None
    """
    if len(doc.tables) == 0:
        print("The document does not contain any tables.")
    else:
        print(f"The document contains {len(doc.tables)} tables.")
        for i, table in enumerate(doc.tables):
            print(f"\nTable ({i}):")
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text)

    return None


def list_sections(doc: Document) -> None:
    """
    Print the sections of a Word document to the console.

    Parameters:
    - doc (Document): The Document object.

    Returns:
    - None
    """
    if len(doc.sections) == 0:
        print("The document does not contain any sections.")
    else:
        print(f"The document contains {len(doc.sections)} sections.")
        # print(dir(doc.sections))
        for i, section in enumerate(doc.sections):
            print(f"\nSection ({i}):")
            print(f"Section start: {section.start_type}")

    return None


def list_runs(doc: Document) -> None:
    """
    Print the runs of a Word document to the console.

    Parameters:
    - doc (Document): The Document object.

    Returns:
    - None
    """
    if len(doc.paragraphs) == 0:
        print("The document does not contain any paragraphs.")
    else:
        print(f"The document contains {len(doc.paragraphs)} paragraphs.")
        for i, paragraph in enumerate(doc.paragraphs):
            print(f"\nParagraph ({i}):")
            for j, run in enumerate(paragraph.runs):
                print(f"\tRun ({j}):")
                print(f"\tText: {run.text}")
                print(f"\tBold: {run.bold}")
                print(f"\tItalic: {run.italic}")
                print(f"\tUnderline: {run.underline}")
                print(f"\tFont Name: {run.font.name}")
                print(f"\tFont Size: {run.font.size}")
                print(f"\tFont Color: {run.font.color.rgb}")
                print(f"\tFont Bold: {run.font.bold}")
                print(f"\tFont Italic: {run.font.italic}")
                print(f"\tFont Underline: {run.font.underline}")
                print(f"\tFont Strike: {run.font.strike}")
                print(f"\tFont Subscript: {run.font.subscript}")
                print(f"\tFont Superscript: {run.font.superscript}")
                print(f"\tFont All Caps: {run.font.all_caps}")
                print(f"\tFont Hidden: {run.font.hidden}")
                print(f"\tFont Highlight Color: {run.font.highlight_color}")
                print(f"\tFont Shadow: {run.font.shadow}")
                print(f"\tFont Size: {run.font.size}")

    return None


def list_styles(doc: Document) -> None:
    """
    Print the styles of a Word document to the console.

    Parameters:
    - doc (Document): The Document object.

    Returns:
    - None
    """
    if len(doc.styles) == 0:
        print("The document does not contain any styles.")
    else:
        print(f"The document contains {len(doc.styles)} styles:")
        for style in doc.styles:
            print(style.name)

    return None
