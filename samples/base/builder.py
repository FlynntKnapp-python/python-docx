# samples\base\builder.py

import math
import os
from typing import Any
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


# To make this work effectively, we may need to have a way to either return a paragraph
# or specify the paragraph to work on.
class DocxBuilder:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = Document()

    def load_or_create(self):
        if os.path.exists(self.file_path):
            print(f"Loading {self.file_path}...")
            self.document = Document(self.file_path)
        else:
            print(f"The file {self.file_path} does not exist. Creating a new one...")
            self.document = Document()

    def save(self):
        """
        Save the document to the file path specified in the constructor.

        Returns:
            bool: True if the document was saved successfully, False otherwise.
            Any: The document object if it was saved successfully, None otherwise.
        """
        try:
            print(f"Saving the document to {self.file_path}...")
            self.document.save(self.file_path)
            print(f"Document saved to {self.file_path}.")
            return True, self.document
        except Exception as e:
            print(f"Error saving the document: {e}")
            return False, None

    def delete_and_create(self):
        """
        Delete the existing file at the specified path and create a new document.

        Returns:
            Document: The newly created document.
        """
        if os.path.exists(self.file_path):
            print(f"Deleting {self.file_path}...")
            os.remove(self.file_path)
        print(f"Creating {self.file_path}...")
        self.document = Document()
        return self.document

    def add_table(self, items: list, cols: int):
        """
        Add a table to the document with the given items and number of columns.

        Args:
            items (list): The items to add to the table.
            cols (int): The number of columns in the table.

        Returns:
            Table: The table object that was added to the document.
        """
        rows = (len(items) + cols - 1) // cols
        table = self.document.add_table(rows=rows, cols=cols)
        for i, item in enumerate(items):
            row, col = divmod(i, cols)
            table.cell(row, col).text = str(item)
            for paragraph in table.cell(row, col).paragraphs:
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
        return table

    # Add other methods here by translating existing functions to methods.
    # Methods that modify the document should use `self.document` instead of passing `doc` as a parameter.
    # For methods that only need to read or perform actions that don't modify the document, you might still pass `Document` objects around if necessary.

    # Example method adaptation for add_resume_heading
    def add_resume_heading(self, name: str, title: str):
        """
        Add a resume heading to the document with the given name and title.

        Args:
            name (str): The name to display in the heading.
            title (str): The title to display in the heading.

        Returns:
            Tuple[Paragraph, Paragraph]: A tuple containing the name and title paragraphs.
        """
        name_paragraph = self.document.add_paragraph()
        name_paragraph.paragraph_format.space_after = 0
        name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_run = name_paragraph.add_run(name)
        name_run.font.name = "Times New Roman"
        name_run.font.size = Pt(18)
        name_run.bold = True

        title_paragraph = self.document.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(12)

        return name_paragraph, title_paragraph

    # You should convert all other functions in a similar manner,
    # ensuring they operate on `self.document` and make use of `self.file_path` where appropriate.
