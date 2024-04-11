# samples\horizontal_line_example.py

from base import docx_builder
from docx import Document

# Specify the file path for the .docx file:
file_path = "samples/output/HorizontalLineExample.docx"

# Create a new Document:
doc = Document()

# Add a paragraph that will contain the horizontal line:
paragraph = doc.add_paragraph()
paragraph.add_run("This is the text above the first line.")

# Add a horizontal line to the bottom of the paragraph:
paragraph = docx_builder.insert_horizontal_line_paragraph_bottom(paragraph)

# Add another paragraph after the line:
paragraph_after = doc.add_paragraph("This is the text below the first line.")

# Add a stand-alone horizontal line:
doc = docx_builder.insert_horizontal_line(doc)

# Add another paragraph after the line:
paragraph_after = doc.add_paragraph("This is the text below the second line.")

# Save the document:
saved = docx_builder.manage_docx_file(file_path, doc, "save")
print("Saved: ", saved)
