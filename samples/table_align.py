# samples\table_align.py

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import utils

# Specify the file path for the .docx file
file_path = "samples/output/TableAlign.docx"

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading("Table Alignment", 0)

# Create a 1-row, 3-column table
table = doc.add_table(rows=1, cols=3)

# Set the text for each cell
table.cell(0, 0).text = "Left\n\nL: Line 2"
table.cell(0, 1).text = "Center\n\nC: Line 2"
table.cell(0, 2).text = "Right\n\nR: Line 2"

# Add a paragraph and align left cell to the left
left_cell_paragraph = table.cell(0, 0).paragraphs[0]
left_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Add a paragraph and align center cell to the center
center_cell_paragraph = table.cell(0, 1).paragraphs[0]
center_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add a paragraph and align right cell to the right
right_cell_paragraph = table.cell(0, 2).paragraphs[0]
right_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Save the document to a .docx file
saved = utils.save_docx(file_path, doc)
