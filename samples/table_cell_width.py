# samples\table_align.py
import os

from base import docx_builder
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# Specify the file path for the .docx file
file_path = "samples/output/TableCellWidth.docx"


# Create a new Document
# doc = Document(file_path)
doc = Document()

# Set the document margins:
doc = docx_builder.set_margins(doc, 0.75, 0.75, 0.75, 0.75)

# TODO Fix the use of this "magic number" in the code below:
# "9" seems to work best for the page width in inches.
# Specify the page width in inches:
page_width_in_inches = 9

# Add a title to the document
doc.add_heading(f"Table Alignment (Width: {page_width_in_inches})", 0)

# Calculate page width:
# page_width_in_inches = docx_builder.get_page_width_in_inches(doc, 0)
# TODO: Why does 10 work but not 8.5?
# TODO: Also, fix "get_page_width_in_inches".

# Calculate column widths in inches:
left_column_width = int(page_width_in_inches * 1 / 3)
right_column_width = int(page_width_in_inches * 2 / 3)


# Create a 1-row, 2-column table
table = doc.add_table(
    rows=3,
    cols=2,
)

###############################
# This does not work as needed. Need to set the individual cell widths.
# Set the width of the columns
# table.columns[0].width = left_column_width  # Left column width
# table.columns[1].width = right_column_width  # Right column width
###############################

# Set the width of the columns
for row in table.rows:
    row.cells[0].width = Inches(left_column_width)  # Left column width
    row.cells[1].width = Inches(right_column_width)  # Right column width

# # Set the width of the columns
# for row in table.rows:
#     row.cells[0].width = left_column_width  # Left column width
#     row.cells[1].width = right_column_width  # Right column width

# Get name and title from the environment:
name = os.getenv("NAME", "John Doe")
title = os.getenv("TITLE", "Software Engineer")

# Get social links from the environment:
email = os.getenv("EMAIL", "john.doe@example.com")
github = os.getenv("GITHUB", "github.com/johndoe")
linkedin = os.getenv("LINKEDIN", "linkedin.com/in/johndoe")

# Set the text for user name and title:
table.cell(0, 0).text = name
table.cell(1, 0).text = title

# Set the text for social links:
table.cell(0, 1).text = f"Email: {email}"
table.cell(1, 1).text = f"GitHub: {github}"
table.cell(2, 1).text = f"LinkedIn: {linkedin}"


# Add a paragraph and align name cell to the left:
left_cell_paragraph = table.cell(0, 0).paragraphs[0]
left_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Add a paragraph and align title cell to the left:
left_cell_paragraph = table.cell(1, 0).paragraphs[0]
left_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Add a paragraph and align email cell to the right:
left_cell_paragraph = table.cell(0, 1).paragraphs[0]
left_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Add a paragraph and align github cell to the right:
left_cell_paragraph = table.cell(1, 1).paragraphs[0]
left_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Add a paragraph and align linkedin cell to the right:
left_cell_paragraph = table.cell(2, 1).paragraphs[0]
left_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Save the document to a .docx file
saved = docx_builder.manage_docx_file(file_path, doc, "save")
print(f"Saved ({saved}):  {file_path}")
