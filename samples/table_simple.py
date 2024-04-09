# samples\table_simple.py

from base import docx_builder
from docx import Document

# Specify the file path for the .docx file:
file_path = "samples/output/TableSimple.docx"

# Create a new Document:
doc = Document()

# Add a title to the document:
doc.add_heading("The Official Table Example!", 0)

# Create a list of records:
records = (
    ("Skill 00", "Skill 10", "Skill 20"),
    ("Skill 01", "Skill 11", "Skill 21"),
    ("Skill 02", "Skill 12", "Skill 22"),
)

# Add a table, which contains the records, to the document:
table = doc.add_table(rows=1, cols=3)
for col0, col1, col2 in records:
    row_cells = table.add_row().cells
    row_cells[0].text = col0
    row_cells[1].text = col1
    row_cells[2].text = col2


# Save the document to a .docx file:
saved = docx_builder.save_docx(file_path, doc)
