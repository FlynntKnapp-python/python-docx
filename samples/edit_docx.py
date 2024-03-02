from docx import Document

from utils import load_docx_if_exists, create_docx_if_not_exists, delete_and_save_docx

file_path = "samples/output/Edit.docx"

# Load the document from a .docx file
doc = load_docx_if_exists(file_path)

if len(doc.paragraphs) == 0:
    print("The document does not contain any paragraphs.")
else:
    print(f"The document contains {len(doc.paragraphs)} paragraphs.")
    for paragraph in doc.paragraphs:
        print("\nParagraph:")
        print(paragraph.text)

# Add a paragraph of text
p = doc.add_paragraph(
    "This is a simple paragraph that is being added to the document. "
)

# Add another heading to the document
doc.add_heading(f"Another Heading: {len(doc.paragraphs)}", level=1)

# Create the .docx file if it does not exist
create_docx_if_not_exists(file_path, doc)

# Delete the file if it exists and save the document to a .docx file
delete_and_save_docx(file_path, doc)
