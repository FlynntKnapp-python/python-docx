# samples\utils.py

import os
from typing import Any

from docx import Document


def manage_docx_file(
    path: str, action: str = "load_or_create", document: Document = None
) -> Document:
    """
    Manage a .docx file (load, create, or delete and recreate).

    Parameters:
    - path (str): The file system path where the .docx file is saved.
    - action (str): Action to perform on the file: 'load_or_create', 'delete_and_create', or 'delete_and_save'.
    - document (Document): The docx.Document object. Required for 'delete_and_save' action.

    Returns:
    - Document: For 'load_or_create' and 'delete_and_create' actions, returns the Document object.
    """
    file_exists = os.path.exists(path)

    if action == "load_or_create":
        if file_exists:
            print(f"Loading {path}...")
            return Document(path)
        else:
            print(f"The file {path} does not exist. Creating a new one...")
            return Document()

    if action == "delete_and_create":
        if file_exists:
            print(f"Deleting {path}...")
            os.remove(path)
        print(f"Creating {path}...")
        return Document()

    if action == "delete_and_save":
        if file_exists:
            print(f"Deleting {path}...")
            os.remove(path)
        if document:
            document.save(path)
            print(f"Document saved to {path}.")


def add_table(doc: Document, items: list, cols: int) -> Document:
    """
    Create and populate a table in the document.

    Parameters:
    - doc (Document): The Document object to add the table to.
    - items (list): List of items to populate the table.
    - cols (int): Number of columns in the table.

    Returns:
    - Document: The modified Document object.
    """
    rows = (len(items) + cols - 1) // cols  # Calculate rows needed
    table = doc.add_table(rows=rows, cols=cols)

    for i, item in enumerate(items):
        row, col = divmod(i, cols)
        table.cell(row, col).text = str(item)

    return doc


def save_docx(path: str, document: Document):
    """
    Save a document to a .docx file.

    Parameters:
    - path (str): The file system path where the .docx file is saved.
    - document (Document): The docx.Document object to be saved.

    Returns:
    - bool: True if the document was saved successfully, False otherwise.
    """
    try:
        print(f"Saving the document to {path}...")
        document.save(path)
        print(f"Document saved to {path}.")
        return True
    except Exception as e:
        print(f"Error saving the document: {e}")
        return False


def enumerate_paragraphs(doc: Document) -> list:
    """
    Returns an enumerated list of paragraphs in a Word document.

    Parameters:
    - doc (Document): The Document object.

    Returns:
    - list: The enumerated list of paragraphs in the document.
    """
    return [(i, paragraph.text) for i, paragraph in enumerate(doc.paragraphs)]


def print_attributes(obj: Any):
    """
    Print the attributes of an object to the console.

    Parameters:
    - obj: The object for which to list attributes.

    Returns:
    - None
    """
    print(f"List of {type(obj).__name__} attributes:")
    for attr in dir(obj):
        print(f"\t{attr}")

    return None
