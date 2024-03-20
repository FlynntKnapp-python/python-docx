from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from utils import delete_and_or_save_docx, create_docx_if_not_exists

# Specify the file path for the .docx file
file_path = "samples/output/TableAlignLeft.docx"

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading("Table Alignment", 0)

# Create a 1-row, 3-column table
table = doc.add_table(rows=1, cols=3)

# Set the text for each cell
table.cell(0, 0).text = "Left"
table.cell(0, 1).text = "Middle"
table.cell(0, 2).text = "Right"

# Add a pragraph and align left cell to the left
left_cell_paragraph = table.cell(0, 0).paragraphs[0]
left_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Add a pragraph and align center cell to the center
center_cell_paragraph = table.cell(0, 1).paragraphs[0]
center_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Add a pragraph and align right cell to the right
right_cell_paragraph = table.cell(0, 2).paragraphs[0]
right_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Delete the file if it exists and save the document to a .docx file
delete_and_or_save_docx(file_path, doc)
