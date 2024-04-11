# samples\demo_docx.py

from base import docx_builder
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

# Specify the file path for the .docx file:
file_path = "samples/output/Demo.docx"

# Create a new Document:
document = Document()

# The `heading` is added as a paragraph with a style of `Heading 0`:
document.add_heading("Document Title", 0)  # Added as paragraph

p = document.add_paragraph("A special paragraph having")  # Added as paragraph
p.add_run("bold").bold = True
p.add_run(", ")
p.add_run("orange text").font.color.rgb = RGBColor(255, 165, 0)
p.add_run(", ")
p.add_run("blue text").font.color.rgb = RGBColor(0, 0, 255)
p.add_run(", ")
p.add_run("burgundy text").font.color.rgb = RGBColor(128, 0, 32)
p.add_run(", ")
p.add_run("20 pt").font.size = Pt(20)
p.add_run(", ")
p.add_run("underline").underline = True
p.add_run(", ")
p.add_run("superscript").font.superscript = True
p.add_run(", ")
p.add_run("subscript").font.subscript = True
p.add_run(", ")
p.add_run("highlighted").font.highlight_color = WD_COLOR_INDEX.YELLOW
p.add_run(", ")
p.add_run("shadowed").font.shadow = True
p.add_run(", ")
p.add_run("all caps").font.all_caps = True
p.add_run(", ")
p.add_run("hidden").font.hidden = True
p.add_run(", ")
p.add_run("comicsans").font.name = "Comic Sans MS"
p.add_run(", ")
p.add_run("Heading 1 Char", style="Heading 1 Char")
p.add_run(", ")
# Alternate way to add a run with a style:
p.add_run("Heading 2 Char").style = "Heading 2 Char"
p.add_run(", ")
p.add_run("Heading 3 Char").style = "Heading 3 Char"
p.add_run(", ")
p.add_run("strike").font.strike = True
p.add_run(", and ")
p.add_run("italic").italic = True
p.add_run(".")

# Heading level can be an arg `1` or a kwarg `level=1`:
document.add_heading("Heading - level 1", level=1)  # Added as paragraph
document.add_paragraph(
    "A specific 'Intense Quote'?", style="Intense Quote"
)  # Added as paragraph
document.add_paragraph("A specific 'Quote'!", style="Quote")  # Added as paragraph

document.add_paragraph(
    "First item in unordered list", style="List Bullet"
)  # Added as paragraph
document.add_paragraph(
    "First item in ordered list", style="List Number"
)  # Added as paragraph
document.add_paragraph(
    "Second item in ordered list", style="List Number"
)  # Added as paragraph
document.add_paragraph(
    "1st item in ordered list", style="List Number 2"
)  # Added as paragraph
document.add_paragraph(
    "2nd item in ordered list", style="List Number 2"
)  # Added as paragraph
document.add_paragraph(
    "Fourth item in ordered list", style="List Number 3"
)  # Added as paragraph
document.add_paragraph()  # Added as paragraph

document.add_paragraph("This is a picture: ")  # Added as paragraph

document.add_picture(
    "images/NoImageAvailable.png", width=Inches(1.25)
)  # Added as paragraph

records = (
    (3, "101", "Spam"),
    (7, "422", "Eggs"),
    (4, "631", "Spam, spam, eggs, and spam"),
)

# Create a 1-row, 3-column table for the header labels:
table = document.add_table(rows=1, cols=3)  # The table is not added as a paragraph
table_hdr_cells = table.rows[0].cells
table_hdr_cells[0].text = "Qty"
table_hdr_cells[1].text = "Id"
table_hdr_cells[2].text = "Desc"

# Set the font size for the header cells:
for cell in table_hdr_cells:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(16)

# Add the records to the table:
for qty, id, desc in records:
    print(f"We have {qty} of item {id}: {desc}")
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()  # Added as paragraph

ages_and_colors = (
    (23, "Orange", "Bob"),
    (42, "Apple", "Alice"),
    (37, "Banana", "Charlie"),
    (11, "Pear", "David"),
)
ages_and_colors_table = document.add_table(
    rows=1, cols=3
)  # The table is not added as a paragraph
ages_and_colors_table.style = "Table Grid"
a_c_hdr_cells = ages_and_colors_table.rows[0].cells
a_c_hdr_cells[0].text = "Age"
a_c_hdr_cells[1].text = "Favorite Color"
a_c_hdr_cells[2].text = "Name"

# Set the font size for the header cells:
for cell in a_c_hdr_cells:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(20)

# Add the records to the table:
for age, color, name in ages_and_colors:
    print(f"{name} is {age} years old and likes {color}(s)")
    row_cells = ages_and_colors_table.add_row().cells
    row_cells[0].text = str(age)
    row_cells[1].text = color
    row_cells[2].text = name

# Add a page break:
document.add_page_break()  # Added as paragraph


# Add header `Second Page` to the second page:
document.add_heading("Second Page", level=0)  # Added as paragraph

# Add another page break:
document.add_page_break()  # Added as paragraph

# Add `This page left intentionally blank.` to the third page:
third_page_lone_paragraph = document.add_paragraph(
    "This page left intentionally blank."
)  # Added as paragraph
third_page_lone_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Save the document to a .docx file
saved = docx_builder.save_docx(file_path, document)
