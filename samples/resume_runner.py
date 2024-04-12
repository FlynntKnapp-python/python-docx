# samples\resume_runner.py

from base import docx_builder
from docx import Document
import os

# Specify the file path for the .docx file:
file_path = "samples/output/ResumeRunner.docx"

# Create a new Document:
doc = Document()

# Set the document margins:
doc = docx_builder.set_margins(doc, 0.75, 0.75, 0.75, 0.75)

# Get name and title from the environment:
name = os.getenv("NAME")
title = os.getenv("TITLE")

# Add name and title to the document:
doc = docx_builder.add_resume_heading(doc, name, title)

# Get address, city, state, and zip code from the environment:
address = os.getenv("ADDRESS")
city = os.getenv("CITY")
state = os.getenv("STATE")
zip = os.getenv("ZIP")

# Add the address to the document:
doc = docx_builder.add_resume_address(doc, address, city, state, zip)

# Get phone number from the environment:
phone = os.getenv("PHONE")

# Add phone number to the document:
doc = docx_builder.add_resume_phone(doc, phone)

# Get email address, github, and linkedin from the environment:
email = os.getenv("EMAIL")
github = os.getenv("GITHUB")
linkedin = os.getenv("LINKEDIN")

# Add email, github, and linkedin to the document:
links = {"Email": email, "GitHub": github, "LinkedIn": linkedin}
doc = docx_builder.add_resume_personal_links(doc, links)

# Specify summary text and add it to the document:
summary_text = "Buckeye Ipsum woody Archie Griffin WOSU oval mirror lake horseshoe moritz scarlet gray Brutus buckeye leaf ohio state the union tbdbitl OH-IO script ohio carmen ohio Hagerty hall Fisher John Glenn buckeyes Lee Horvath Michael Redd the Lantern Morrill Hayes st. john's excellence Urban Meyer Hoppalong Cassidy"  # noqa E501
summary = doc.add_paragraph(summary_text)
summary = docx_builder.insert_horizontal_line_paragraph_top(summary)
summary = docx_builder.insert_horizontal_line_paragraph_bottom(summary)

# Define a list of skills:
skills = [
    "Git",
    "Scrum",
    "Agile",
    "Python",
    "Django",
    "Django REST",
    "Docker",
    "Linux",
    "S3",
    "Raspberry Pi",
    "Raspberry Pi Pico",
]

# Add a skills table:
table_title = doc.add_paragraph("Skills:")
table_title.style = "Heading 1"
for run in table_title.runs:
    run.underline = True

doc = docx_builder.add_table(doc, skills, 3)

# Add a horizontal line:
doc = docx_builder.insert_horizontal_line(doc)

# Save the document to a .docx file:
saved = docx_builder.manage_docx_file(file_path, doc, "save")
print("Saved: ", saved)
