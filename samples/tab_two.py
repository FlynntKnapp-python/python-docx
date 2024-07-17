from docx import Document
from docx.shared import Inches

file_path = "samples/output/TabTwo.docx"


# Function to convert a list of string items to a single string with tab stops inserted
# between the items:
def add_tab_breaks(items):
    pass


# Create a new Document
doc = Document()

# Add a new paragraph
p = doc.add_paragraph()

# Define the tab stops
tab_stops = p.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(1.5))
tab_stops.add_tab_stop(Inches(3.0))
tab_stops.add_tab_stop(Inches(4.5))

# Add the text with tab stops
run = p.add_run("One\tTwo\tThree\tFour")

# Save the document
doc.save(file_path)
