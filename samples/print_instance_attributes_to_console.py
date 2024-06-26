# samples\print_instance_attributes_to_console.py

from base import docx_builder
from docx import Document

# Create a new (empty) Document:
doc = Document()

# List the attributes of the Document object:
attrs = docx_builder.print_attributes(doc)

# Output:
"""powershell
List of <docx.document.Document object at 0x0000028F9B39E3F0> Attributes:
        _Document__body
        __class__
        __delattr__
        __dict__
        __dir__
        __doc__
        __eq__
        __format__
        __ge__
        __getattribute__
        __getstate__
        __gt__
        __hash__
        __init__
        __init_subclass__
        __le__
        __lt__
        __module__
        __ne__
        __new__
        __reduce__
        __reduce_ex__
        __repr__
        __setattr__
        __sizeof__
        __str__
        __subclasshook__
        __weakref__
        _block_width
        _body
        _element
        _parent
        _part
        add_heading
        add_page_break
        add_paragraph
        add_picture
        add_section
        add_table
        core_properties
        element
        inline_shapes
        iter_inner_content
        paragraphs
        part
        save
        sections
        settings
        styles
        tables
"""
