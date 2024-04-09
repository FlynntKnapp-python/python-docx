# samples\two_runs_in_paragraph.py

from pprint import pprint

from base import docx_builder
from docx.shared import Pt

# Specify the file path for the .docx file:
file_path = "samples/output/Two-Runs-In-Paragraph.docx"

# Load the .docx file (instantiating the Document object):
doc = docx_builder.delete_and_create_docx(file_path)


# Add a paragraph to the document:
paragraph = doc.add_paragraph()
print("dir(paragraph):")
pprint(dir(paragraph))
print("paragraph.text: ", paragraph.text)

# Add a run to the paragraph:
run1 = paragraph.add_run("This is the first run.")
run1.font.size = Pt(12)
print("dir(run1):")
pprint(dir(run1))
print("run1.text: ", run1.text)

# Add a run to the paragraph:
run2 = paragraph.add_run(" This is the second run.")
run2.font.size = Pt(18)
print("run2.text: ", run2.text)


docx_builder.save_docx(file_path, doc)
