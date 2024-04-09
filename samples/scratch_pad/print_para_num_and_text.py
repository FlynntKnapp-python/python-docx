# samples\print_para_num_and_text.py

from docx import Document


# Specify the file path for the .docx file:
file_path = "samples/output/ParaTablePara.docx"

# Load the .docx file (instantiating the Document object):
doc = Document(file_path)

# Get the number of paragraphs in the document:
num_paragraphs = len(doc.paragraphs)

# Print the number of paragraphs in the document:
print(f"Number of Paragraphs: {num_paragraphs}")

# Print the text of each paragraph in the document:
for i, paragraph in enumerate(doc.paragraphs):
    print(f"Paragraph {i}: {paragraph.text}")
