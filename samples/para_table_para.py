# samples\para_table_para.py

from base import docx_builder
from docx import Document

# Specify the file path for the .docx file
file_path = "samples/output/ParaTablePara.docx"

paragraph_text_01 = (
    "This is the first paragraph of text that is being added to the document."
)
items = [
    "Grits",
    "Gravy",
    "Biscuits",
    "Wallets",
    "Keys",
    "CPAP",
    "Glasses",
    "Phone",
    "Kitten",
    "Pupper",
    "Laptop",
]

paragraph_text_02 = (
    "This is the second paragraph of text that is being added to the document."
)


def add_empty_paragraph(doc):
    """
    Add an empty paragraph to the document.

    Args:
        doc: The Document object to which the empty paragraph will be added.

    Returns:
        The Document object with the empty paragraph added.
    """
    doc.add_paragraph()
    return doc


# Create a new Document:
doc = Document()

# Add a paragraph of text:
p1 = doc.add_paragraph(paragraph_text_01)

# Add an empty paragraph:
add_empty_paragraph(doc)

# Add a Table to the Document:
doc = docx_builder.add_table(doc, items, cols=4)

# Add an empty paragraph:
doc = add_empty_paragraph(doc)

# Add a paragraph of text:
p2 = doc.add_paragraph(paragraph_text_02)

# Print the number of paragraphs in the document:
print(f"Number of paragraphs in the document: {len(doc.paragraphs)}")

# Save the document to a .docx file:
saved = docx_builder.manage_docx_file(file_path, doc, "save")
