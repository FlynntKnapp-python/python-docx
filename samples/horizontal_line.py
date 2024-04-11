# samples\horizontal_line.py

from base import docx_builder
from docx import Document
from docx.text.paragraph import Paragraph

file_path = "samples/output/HorizontalLine.docx"

# Create a new Document:
doc = Document()

# Add a paragraph to the document:
p1 = doc.add_paragraph("This is a paragraph.")

# Add a horizontal line:
doc = docx_builder.insert_horizontal_line(doc)

# Add another paragraph to the document:
p2 = doc.add_paragraph("This is another paragraph.")

# Save the document:
saved = docx_builder.save_docx(file_path, doc)
print("Saved: ", saved)
