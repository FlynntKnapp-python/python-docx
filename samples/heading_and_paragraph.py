# samples\heading_and_paragraph.py

from base import docx_builder
from docx import Document

# Specify the file path for the .docx file:
file_path = "samples/output/HeadingAndParagraph.docx"

# Create a new Document:
doc = Document()

# Add a title to the document:
h = doc.add_heading("Document Title", 0)
print("Added h.text: \n", h.text)

# Add a paragraph of text:
p = doc.add_paragraph(
    "This is a simple paragraph that is being added to the document. "
)
print("Added p.text: \n", p.text)

# Save the document to a .docx file:
saved = docx_builder.manage_docx_file(file_path, doc, "save")
print("Saved: ", saved)
