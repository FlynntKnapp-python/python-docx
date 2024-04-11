# samples\horizontal_line_insert_later.py

from base import docx_builder
from docx import Document

# Specify the file path for the .docx file:
file_path = "samples/output/HorizontalLineLater.docx"

# Create a new Document:
doc = Document()

# Add an paragraph:
paragraph = doc.add_paragraph("This is the text above the horizontal line.")

# Save the document:
saved = docx_builder.manage_docx_file(file_path, doc, "save")
print("Saved: ", saved)

# Load the Document:
doc = docx_builder.manage_docx_file(file_path, "load_or_create")

# Get the first paragraph of the document:
paragraph = doc.paragraphs[0]

# Add a horizontal line to the paragraph:
paragraph = docx_builder.insert_horizontal_line_paragraph_bottom(paragraph)

# Save the document:
saved = docx_builder.manage_docx_file(file_path, doc, "save")
print("Saved: ", saved)
